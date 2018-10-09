#coding:utf-8
import win32com
import time
import os
import re

from docx import Document

'''
已能实现循环读入docx文件，并将信息写入同一excel中,初步实现列表查找，比对病人与医生信息，将信息填入合适位置。
'''

#将doc文件另存为docx文件
def doc_to_docx(path):
    if os.path.splitext(path)[1] == ".doc":
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(path)  #目标路径下的文件
        doc.SaveAs(os.path.splitext(path)[0]+".docx",16)  #转化后路径下的文件,16代表另存为docx文件
        doc.Close()
        word.Quit()
        
#获取文件夹下的所有文件的绝对路径
def find_file(path, ext, file_list=[]):
    dir = os.listdir(path)
    for i in dir:
        i = os.path.join(path, i)
        if os.path.isdir(i):
            find_file(i, ext, file_list)
        else:
            if ext == os.path.splitext(i)[1]:
                file_list.append(i)
    return file_list

'''将word表格信息写入excel'''
def write_excel(docx_path,i_in,workbook,name,doctor,seq_num,temputure,mailv,hupin,bld_presure_left,
                bld_presure_left2,bld_presure_right,bld_presure_right2,body_hig,body_weig,
                yaowei,BMI,self_estimate_num,self_ability_num,self_cognition_num,
                self_motion_num,exer_fre_num,exer_time,exer_holdon,exer_mode,eat_habit_num,
                smok_condt_num,smok_perday_num,smok_beg_age,drink_fre,sight_left,sight_right,
                hear,exr_ability,skin,gongmo,linba,tzx,hxy,luoyin,zbdm,kfxt,
                xdt,nwl,thhdb,xbx,B_super,nxgb,szb,xzb,xgb,ybb,sjb,qtb,jkzd
                ):
      
    sheet=workbook.Worksheets(1) #获取当前第一个表格
    nrows = sheet.UsedRange.Rows.Count #获取excel表格当前行数
    print(nrows)
    #匹配word文件中的姓名和医生信息，将信息填入相应表格中
    #如果找不到，向屏幕打印找不到此人
    for num in range(2,nrows):
        if str(sheet.Cells(num,'A').text.split()) == str(name) and str(sheet.Cells(num,'K').text.split()) == str(doctor):
            i_in = num
            break
        elif num == (nrows-1) and str(sheet.Cells(num,'A').text.split()) != str(name) :
            sheet.Rows(nrows+1).Insert()
            i_in = nrows+1
            sheet.Cells(i_in,'A').value = name
            sheet.Cells(i_in,'K').value = doctor
            print('Can not find this people, please check!')
            break
    sheet.Cells(i_in,'M').value = seq_num
    sheet.Cells(i_in,'N').value = temputure
    sheet.Cells(i_in,'O').value = mailv
    sheet.Cells(i_in,'P').value = hupin
    sheet.Cells(i_in,'Q').value = bld_presure_left
    sheet.Cells(i_in,'R').value = bld_presure_left2
    sheet.Cells(i_in,'S').value = bld_presure_right
    sheet.Cells(i_in,'T').value = bld_presure_right2
    sheet.Cells(i_in,'U').value = body_hig
    sheet.Cells(i_in,'V').value = body_weig
    sheet.Cells(i_in,'W').value = yaowei
    sheet.Cells(i_in,'X').value = BMI
    
    sheet.Cells(i_in,'Y').value = self_estimate_num
    sheet.Cells(i_in,'Z').value = self_ability_num
    sheet.Cells(i_in,'AA').value = self_cognition_num
    sheet.Cells(i_in,'AB').value = self_motion_num
    
    sheet.Cells(i_in,'AC').value = exer_fre_num
    sheet.Cells(i_in,'AD').value = exer_time
    sheet.Cells(i_in,'AE').value = exer_holdon
    sheet.Cells(i_in,'AF').value = exer_mode
    sheet.Cells(i_in,'AG').value = eat_habit_num

    sheet.Cells(i_in,'AH').value = smok_condt_num
    sheet.Cells(i_in,'AI').value = smok_perday_num
    sheet.Cells(i_in,'AJ').value = smok_beg_age
    sheet.Cells(i_in,'AL').value = drink_fre
    
    #无法提取sheet.Cells(i_in,'AR').value = smok_perday_num
    #无法提取sheet.Cells(i_in,'AS').value = smok_perday_num
    #无法提取sheet.Cells(i_in,'AT').value = smok_perday_num
    #无法提取sheet.Cells(i_in,'AU').value = smok_perday_num
    sheet.Cells(i_in,'AV').value = sight_left
    sheet.Cells(i_in,'AW').value = sight_right
    sheet.Cells(i_in,'AX').value = hear
    sheet.Cells(i_in,'AY').value = exr_ability

    sheet.Cells(i_in,'BA').value = skin
    sheet.Cells(i_in,'BB').value = gongmo
    sheet.Cells(i_in,'BC').value = linba
    sheet.Cells(i_in,'BD').value = tzx
    sheet.Cells(i_in,'BE').value = hxy
    sheet.Cells(i_in,'BF').value = luoyin
    #无法提取sheet.Cells(i_in,'BG').value = xinlv
    #无法提取sheet.Cells(i_in,'BH').value = xinlv2
    #无法提取sheet.Cells(i_in,'BI').value = zayin
    #无法提取sheet.Cells(i_in,'BJ').value = fubu
    #无法提取sheet.Cells(i_in,'BK').value = fubu
    #无法提取sheet.Cells(i_in,'BL').value = fubu
    #无法提取sheet.Cells(i_in,'BM').value = fubu
    #无法提取sheet.Cells(i_in,'BN').value = fubu

    #sheet.Cells(i_in,'B0').value = xzsz
    sheet.Cells(i_in,'BP').value = zbdm
    #无法提取first_sheet.Cells(i_in,'BT').value = xcg
    sheet.Cells(i_in,'CB').value = kfxt
    sheet.Cells(i_in,'CC').value = xdt
    sheet.Cells(i_in,'CD').value = nwl
    sheet.Cells(i_in,'CF').value = thhdb
    sheet.Cells(i_in,'CU').value = xbx
    sheet.Cells(i_in,'CV').value = B_super

    sheet.Cells(i_in,'CW').value = nxgb
    sheet.Cells(i_in,'CX').value = szb
    sheet.Cells(i_in,'CY').value = xzb
    sheet.Cells(i_in,'CZ').value = xgb
    sheet.Cells(i_in,'DA').value = ybb
    sheet.Cells(i_in,'DB').value = sjb
    sheet.Cells(i_in,'DC').value = qtb
    sheet.Cells(i_in,'DR').value = jkzd
    workbook.Save()

'''
从word文档中提取所需信息
'''
#读取docx，返回所需信息
def parse_docx(docx_path,i_in,workbook):
    document = Document(docx_path) #读入文件,doc_path为文档的路径
    tables = document.tables #获取文件中的表格集,word文档中可能有多个表格

    ''' 第一张表格 '''
    table = tables[0]  #通过下标,获取文件中的第一个表格
    #title= document.paragraphs[0].text #文档标题
    name = table.cell(0,1).text.split()   #患者姓名,<class 'str'>
    #time = table.cell(1,2).text   #体检日期
    doctor = table.cell(1,5).text.split() #责任医生,<class 'str'>
    
    #症状
    seq = table.cell(3,1).text.split() #症状,<class 'list'>
    seq_num = [] #症状代号
    for n_i in range(50,len(seq)):
        if seq[n_i] != '/':
           seq_num.append(seq[n_i])
           ++n_i

    #一般情况
    temputure = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(4,2).text)  #体温,re.findall()提取数字
    mailv = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(4,5).text)  #脉率
    hupin = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(5,2).text) #呼吸频率
    bld_presure_left = table.cell(5,6).text.split()[0] #血压左侧收缩
    bld_presure_left2 = table.cell(5,6).text.split()[2] #血压左侧舒张
    bld_presure_right = table.cell(6,6).text.split()[0] #血压右侧收缩
    bld_presure_right2 = table.cell(6,6).text.split()[2] #血压右侧舒张 
    body_hig = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(7,2).text) #身高
    body_weig = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(7,5).text) #体重
    yaowei = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(8,2).text) #腰围
    BMI = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(8,5).text) #体质指数
  
    self_estimate = table.cell(9,2).text #老年人健康状态自我评估
    self_estimate_num = self_estimate.split()[-1]
    self_ability = table.cell(10,2).text #老年人自理能力自我评估
    self_ability_num = self_ability.split()[-1]
    self_cognition = table.cell(11,2).text #老年人认知功能
    self_cognition_num = self_cognition.split()[-1]
    self_motion = table.cell(12,2).text #老年人情感状态
    self_motion_num = self_motion.split()[-1]

    #生活方式
    exer_fre = table.cell(13,3).text #锻炼频率
    exer_fre_num = exer_fre.split()[-1]
    exer_time = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(14,3).text) #每次锻炼时间
    exer_holdon = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(14,5).text) #坚持锻炼时间
    exer_mode = table.cell(15,3).text #锻炼方式
    eat_habit = table.cell(16,2).text.split() #饮食习惯
    eat_habit_num = []
    for n_j in range(12,len(eat_habit)):
        if eat_habit[n_j] != '/':
           eat_habit_num.append(eat_habit[n_j])
           ++n_j

    smok_condt_num = table.cell(17,3).text.split()[-1] #吸烟状况
    smok_perday_num = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(18,3).text) #日吸烟量
    smok_beg_age = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(19,3).text) #开始吸烟年龄
    smok_end_age = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(19,5).text) #戒烟年龄
    drink_fre = table.cell(20,3).text.split()[-1] #饮酒频率
    #drink_perday = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(21,3).text) #日饮酒量
  
    ''' 饮酒戒酒状况不知是否需要统计，数字部分下标为7，但若统计的话，易产生越界情况  '''  
    '''#保留饮酒信息模块
    is_dryout_num = table.cell(22,3).text.split()#[-1] #是否戒酒
    if is_dryout_num == '2':
        dryout_age = is_dryout[4]
    else:
        dryout_age = '0'
    drink_beg_age = re.findall(r'-?\d+.?\d*e?-?\d*?',table.cell(23,3).text) #开始饮酒年龄
    is_drunk = table.cell(23,5).text.split()[-1] #一年内是否醉酒
    drink_mode = table.cell(24,3).text #饮酒种类
    #无法识别
    occupational_disease = table.cell(25,2).text #职业病危害因素接触史，文本无法提取
    #无法识别
    '''
    
    ''' 第二张表格 '''

    table1 = tables[1] #获取文件中的第二个表格
   
    #脏器功能
    #无法识别,原因为此表格中嵌套有其他表格，故选择提取的信息为空
    #mouth = table1.cell(0,3).text #口唇,数据无法提取
    #无法识别
    sight_left = table1.cell(1,3).text.split()[1] #左眼视力
    sight_right = table1.cell(1,3).text.split()[3] #右眼视力
    hear = table1.cell(2,3).text.split()[-1] #听力
    exr_ability = table1.cell(3,3).text.split()[-1] #运动功能
  
    #查体
    skin = table1.cell(5,3).text.split()[-1] #皮肤
    gongmo = table1.cell(6,3).text.split()[-1] #巩膜
    linba = table1.cell(7,3).text.split()[-1] #淋巴结
    tzx = table1.cell(8,3).text.split()[-1] #桶状胸
    hxy = table1.cell(9,3).text.split()[-1] #呼吸音
    luoyin = table1.cell(10,3).text.split()[-1] #罗音
  
    #无法识别
    #xinlv = re.findall(r'-?\d+.?\d*e?-?\d*?',table1.cell(11,3).text.split()[1])
    #xinlv2 = table1.cell(11,3).text.split()
    #zayin = table1.cell(11,3).text.split() 
    #fubu = table1.cell(12,3).text.split()
    #无法识别

    #xzsz = table1.cell(13,3).text.split()[8] #下肢水肿
    zbdm = table1.cell(14,3).text.split()[-1] #足背动脉
  
    #辅助检查
    #无法识别   
    #xcg = table1.cell(23 ,3).text.split() #血常规
    #无法识别
  
    kfxt = re.findall(r'-?\d+.?\d*e?-?\d*?',table1.cell(25,3).text)#空腹血糖
    xdt = table1.cell(26,3).text.split()[-1] #心电图


    ''' 第三张表格 '''

    #辅助检查
    table2 = tables[2] #获取文件中的第三个表格
    nwl = re.findall(r'-?\d+.?\d*e?-?\d*?',table2.cell(0,2).text)
    thhdb = re.findall(r'-?\d+.?\d*e?-?\d*?',table2.cell(2,2).text)
    #无法识别
    #ggn = table2.cell(4,1).text #肝功能
    #sgn = table2.cell(5,1).text #肾功能
    #xz = table2.cell(6,5).text #血脂
    #无法识别,可以将下标改下试试
    xbx = table2.cell(7,2).text.split()[-1] #胸部X线片
    B_super = table2.cell(8,2).text.split()[-1] #B超

    #现存主要健康问题
    nxgb = table2.cell(13,2).text.split()[2] #脑血管疾病
    szb = table2.cell(15,2).text.split()[2] #肾脏疾病
    xzb = table2.cell(17,2).text.split()[-6] #心脏疾病
    xgb = table2.cell(18,2).text.split()[-3] #血管疾病
    ybb = table2.cell(20,2).text.split()[-4] #眼部疾病
    sjb = table2.cell(21,2).text.split()[-1] #神经系统疾病
    qtb = table2.cell(22,2).text.split()[-1] #其他系统疾病
  
    ''' 第四张表格 '''
    table3 = tables[3] #获取文件中的第四个表格,小表格信息无法统计
    #主要用药情况
    #yongy = table3.cell(2,4).text#.split()
    jkzd = table3.cell(3,1).text.split()[-3] #健康指导
    #wykz = table3.cell(2,6).text#.split()
    #调用写入excel函数
    write_excel(docx_path,i_in,workbook,name,doctor,seq_num,temputure,mailv,hupin,bld_presure_left,
                bld_presure_left2,bld_presure_right,bld_presure_right2,body_hig,body_weig,
                yaowei,BMI,self_estimate_num,self_ability_num,self_cognition_num,
                self_motion_num,exer_fre_num,exer_time,exer_holdon,exer_mode,eat_habit_num,
                smok_condt_num,smok_perday_num,smok_beg_age,drink_fre,sight_left,sight_right,
                hear,exr_ability,skin,gongmo,linba,tzx,hxy,luoyin,zbdm,kfxt,
                xdt,nwl,thhdb,xbx,B_super,nxgb,szb,xzb,xgb,ybb,sjb,qtb,jkzd)
  
if __name__ == "__main__":
    start = time.time()
    word = win32com.client.Dispatch('Word.Application') #打开word应用程序
    excel = win32com.client.Dispatch('Excel.Application') 
    workbook = excel.Workbooks.Open(r"C:\Users\xxxx\Desktop\ABC\xxxxx.xls") #统计表所在的绝对路径
    #后台运行,不显示,不警告
    excel.Visible = False #表格可见与否
    excel.DisplayAlerts = False #警告信息是否显示
    #设置初始行数
    begLine = 1
    #遍历文件,将doc文件转换为docx文件
    dir_path = r'C:\Users\xxxx\Desktop\ABC' # doc文件的绝对路径
    ext = ".doc"
    file_list = find_file(dir_path,ext)
    for file in file_list:
        doc_to_docx(file)
    #遍历文件,调用读word和写excel,将信息统计起来
    doc_files = os.listdir(dir_path)
    for doc in doc_files:
        if os.path.splitext(doc)[1] == '.docx':
            try:
                parse_docx(dir_path+'\\'+doc,begLine,workbook)
            except Exception as e:
                print(e)
        elif os.path.splitext(doc)[1] == '.doc':
            print("Can not open this type of file!")     
    print('All done!')
    workbook.Close()
    excel.Application.Quit()
    end = time.time()
    dur_time = end - start
    print("This program cost "+ str(dur_time) + "s") #打印总共耗时
    time.sleep(3)
